from flask import Flask, render_template, request, redirect, url_for, flash, send_file, session
from flask_sqlalchemy import SQLAlchemy
from flask_wtf import FlaskForm
from flask_wtf.csrf import CSRFProtect, generate_csrf
from wtforms import StringField, IntegerField, SelectField, SubmitField, TextAreaField, PasswordField, RadioField
from typing import Dict, List, Optional, Tuple, Union, Set
from wtforms.validators import DataRequired, NumberRange, Email, Length, EqualTo, Optional
from wtforms import ValidationError
from werkzeug.security import generate_password_hash, check_password_hash
import os
from datetime import datetime, time, timedelta
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
import io
import re
from collections import namedtuple

app = Flask(__name__)
app.config['SECRET_KEY'] = 'your-secret-key'
app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get('DATABASE_URL', 'sqlite:///timetable.db')
app.config['TEMPLATES_AUTO_RELOAD'] = True
db = SQLAlchemy(app)
# --- Database configuration (Supabase + local fallback) ---
DATABASE_URL = os.environ.get('DATABASE_URL')
LOCAL_DEV = os.environ.get('LOCAL_DEV', '0') == '1' or os.environ.get('FLASK_ENV') == 'development'

if DATABASE_URL:
    # Fix scheme if needed (Supabase often provides postgres://)
    if DATABASE_URL.startswith('postgres://'):
        DATABASE_URL = DATABASE_URL.replace('postgres://', 'postgresql+psycopg2://', 1)
    app.config['SQLALCHEMY_DATABASE_URI'] = DATABASE_URL
else:
    if LOCAL_DEV:
        print("⚙️ Using local SQLite database (development mode)")
        app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///timetable.db'
    else:
        raise RuntimeError("❌ DATABASE_URL not set. Please configure Supabase connection string.")

app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['SQLALCHEMY_ENGINE_OPTIONS'] = {"pool_pre_ping": True}

# Initialize DB safely
try:
    with app.app_context():
        db.create_all()
        print("✅ Database tables created or verified.")
except Exception as e:
    print(f"⚠️ Database initialization error (ignored): {e}")

csrf = CSRFProtect(app)

# Models
class Course(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    faculty_id = db.Column(db.Integer, db.ForeignKey('faculty.id'))
    classroom_id = db.Column(db.Integer, db.ForeignKey('classroom.id'))
    duration = db.Column(db.Integer, default=1)  # hours
    day = db.Column(db.String(10))
    start_time = db.Column(db.Time)
    department = db.Column(db.String(50))  # cse, ece, me, ee
    year = db.Column(db.Integer)  # 1-4
    semester = db.Column(db.Integer)  # 1-8

    faculty = db.relationship('Faculty', backref='course_list')
    classroom = db.relationship('Classroom', backref='course_list')

class Faculty(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    availability = db.Column(db.String(200))  # JSON or simple string
    max_load = db.Column(db.Integer, default=5)
    department = db.Column(db.String(50))  # cse, ece, me, ee
    year = db.Column(db.Integer)  # 1-4
    semester = db.Column(db.Integer)  # 1-8

    timetables = db.relationship('Timetable', backref='faculty_obj')

class Classroom(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    capacity = db.Column(db.Integer, nullable=False)
    type = db.Column(db.String(50))  # smart-classroom, lab, etc.

    timetables = db.relationship('Timetable', backref='classroom_obj')

class Timetable(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    course_id = db.Column(db.Integer, db.ForeignKey('course.id'))
    faculty_id = db.Column(db.Integer, db.ForeignKey('faculty.id'))
    classroom_id = db.Column(db.Integer, db.ForeignKey('classroom.id'))
    day = db.Column(db.String(10))
    start_time = db.Column(db.Time)
    end_time = db.Column(db.Time)
    department = db.Column(db.String(50))
    semester = db.Column(db.Integer)
    generation = db.Column(db.Integer, default=1)

    course = db.relationship('Course', backref='timetables')

# Association table for User-Course enrollments
enrollments = db.Table('enrollments',
    db.Column('user_id', db.Integer, db.ForeignKey('user.id'), primary_key=True),
    db.Column('course_id', db.Integer, db.ForeignKey('course.id'), primary_key=True)
)

class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    full_name = db.Column(db.String(150), nullable=False)
    email = db.Column(db.String(150), unique=True, nullable=False)
    password_hash = db.Column(db.String(256), nullable=False)
    role = db.Column(db.String(50), nullable=False)  # student, faculty, admin
    department = db.Column(db.String(50))  # cse, ece, me, ee (for students)
    year = db.Column(db.Integer)  # 1-4 (for students)
    semester = db.Column(db.Integer)  # 1-8 (for students)

    enrolled_courses = db.relationship('Course', secondary=enrollments, backref='enrolled_students')

    def set_password(self, password):
        self.password_hash = generate_password_hash(password)

    def check_password(self, password):
        return check_password_hash(self.password_hash, password)

# Forms
class RegistrationForm(FlaskForm):
    full_name = StringField('Full Name', validators=[DataRequired()])
    email = StringField('Email', validators=[DataRequired(), Email()])
    password = PasswordField('Password', validators=[DataRequired(), Length(min=6)])
    confirm_password = PasswordField('Confirm Password', validators=[DataRequired(), EqualTo('password')])
    role = RadioField('Role', choices=[('student', 'Student'), ('faculty', 'Faculty'), ('admin', 'Admin')], default='student')
    department = SelectField('Department', choices=[('cse', 'CSE'), ('ece', 'ECE'), ('me', 'ME'), ('ee', 'EE')])
    year = IntegerField('Year', validators=[Optional(), NumberRange(min=1, max=4)])
    semester = IntegerField('Semester', validators=[Optional(), NumberRange(min=1, max=8)])
    submit = SubmitField('Create Account')

    def validate(self, extra_validators=None):
        rv = FlaskForm.validate(self, extra_validators)
        if not rv:
            return False
        if self.role.data == 'student':
            if not self.year.data:
                self.year.errors = list(self.year.errors) + ['Year is required.']
                rv = False
            if not self.semester.data:
                self.semester.errors = list(self.semester.errors) + ['Semester is required.']
                rv = False
        return rv

class LoginForm(FlaskForm):
    email = StringField('Email', validators=[DataRequired(), Email()])
    password = PasswordField('Password', validators=[DataRequired()])
    submit = SubmitField('Sign In')

class ProfileForm(FlaskForm):
    full_name = StringField('Full Name', validators=[DataRequired()])
    password = PasswordField('New Password', validators=[Length(min=6)], description="Leave blank to keep current password")
    confirm_password = PasswordField('Confirm New Password', validators=[EqualTo('password', message='Passwords must match')])
    department = SelectField('Department', choices=[('cse', 'CSE'), ('ece', 'ECE'), ('me', 'ME'), ('ee', 'EE')])
    year = IntegerField('Year', validators=[Optional(), NumberRange(min=1, max=4)])
    semester = IntegerField('Semester', validators=[Optional(), NumberRange(min=1, max=8)])
    submit = SubmitField('Update Profile')

class CourseForm(FlaskForm):
    name = StringField('Course Name', validators=[DataRequired()])
    faculty_id = SelectField('Faculty', coerce=int, validators=[DataRequired()])
    classroom_id = SelectField('Classroom', coerce=int, validators=[DataRequired()])
    duration = IntegerField('Duration (hours)', default=1, validators=[NumberRange(min=1)])
    department = SelectField('Department', choices=[('cse', 'CSE'), ('ece', 'ECE'), ('me', 'ME'), ('ee', 'EE')])
    year = IntegerField('Year', validators=[NumberRange(min=1, max=4)])
    semester = IntegerField('Semester', validators=[NumberRange(min=1, max=8)])
    submit = SubmitField('Add Course')

class EnrollmentForm(FlaskForm):
    course_id = SelectField('Select Course to Enroll', coerce=int, validators=[DataRequired()])
    submit = SubmitField('Enroll in Course')

class FacultyForm(FlaskForm):
    name = StringField('Faculty Name', validators=[DataRequired()])
    availability = TextAreaField('Availability (e.g., Mon-Fri 9:00-17:00 or 10.00 to 16.00 or 10,00 to 16,00)')
    max_load = IntegerField('Max Load (classes)', default=5)
    department = SelectField('Department', choices=[('cse', 'CSE'), ('ece', 'ECE'), ('me', 'ME'), ('ee', 'EE')])
    year = IntegerField('Year', validators=[NumberRange(min=1, max=4)])
    semester = IntegerField('Semester', validators=[NumberRange(min=1, max=8)])
    submit = SubmitField('Add Faculty')

class ClassroomForm(FlaskForm):
    name = StringField('Classroom Name', validators=[DataRequired()])
    capacity = IntegerField('Capacity', validators=[DataRequired()])
    type = SelectField('Type', choices=[('smart-classroom', 'Smart Classroom'), ('lab', 'Lab'), ('seminar', 'Seminar Hall')])
    submit = SubmitField('Add Classroom')

class ConflictFreeScheduler:
    """Intelligent timetable scheduler with comprehensive conflict detection"""

    def __init__(self, courses: List[Course]):
        self.courses = courses

        # Scheduling constraints
        self.slots = [time(10, 0), time(11, 0), time(12, 0), time(14, 0),
                      time(15, 0), time(16, 0)]

        # Load resources
        self.faculties = {f.id: f for f in Faculty.query.all()}
        self.classrooms = {c.id: c for c in Classroom.query.all()}

        # Parse faculty available days
        self.faculty_days = {f.id: self._parse_available_days(f.availability) for f in self.faculties.values()}
        self.days = sorted(set(d for days in self.faculty_days.values() for d in days))

        # Define day order to prioritize earliest days
        day_order = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']
        self.days = sorted(self.days, key=lambda d: day_order.index(d) if d in day_order else 7)

        # Create time_slots as list of (day, start_time) tuples, sorted by day order then time
        self.time_slots = []
        for day in self.days:
            for slot in self.slots:
                self.time_slots.append((day, slot))
        self.time_slots = sorted(self.time_slots, key=lambda s: (day_order.index(s[0]) if s[0] in day_order else 7, s[1]))

        self.current_gen = None

    def _parse_available_days(self, availability: str) -> List[str]:
        """Parse faculty availability string to extract available days."""
        days_map = {
            'mon': 'Monday', 'tue': 'Tuesday', 'wed': 'Wednesday',
            'thu': 'Thursday', 'fri': 'Friday', 'sat': 'Saturday', 'sun': 'Sunday',
            'monday': 'Monday', 'tuesday': 'Tuesday', 'wednesday': 'Wednesday',
            'thursday': 'Thursday', 'friday': 'Friday', 'saturday': 'Saturday', 'sunday': 'Sunday'
        }
        days = []
        for part in availability.lower().replace(',', ' ').split():
            part = part.strip()
            if part in days_map:
                days.append(days_map[part])
        return list(set(days))  # Remove duplicates



    def generate(self):
        """Generate optimized timetable across all departments with greedy scoring"""
        # Delete all old timetables for global regeneration
        Timetable.query.delete()
        db.session.commit()

        all_courses = self.courses
        all_faculties = self.faculties
        all_classrooms = self.classrooms
        all_slots = sorted(self.time_slots, key=lambda s: (s[0], s[1]))

        timetable = []
        faculty_schedule = {f.id: {d: [] for d in self.days} for f in all_faculties.values()}
        classroom_schedule = {c.id: {d: [] for d in self.days} for c in all_classrooms.values()}
        dept_sem_schedule = {}

        def available(course, faculty, classroom, day, slot):
            # Check if faculty is available on this day
            if day not in self.faculty_days[faculty.id]:
                return False
            # For multi-hour courses, check consecutive slots
            duration = course.duration
            if duration > 1:
                slots_needed = []
                for i in range(duration):
                    next_slot = (datetime.combine(datetime.today(), slot) + timedelta(hours=i)).time()
                    if next_slot not in self.slots:
                        return False
                    slots_needed.append(next_slot)
                # Check all consecutive slots for conflicts
                for s in slots_needed:
                    if s in faculty_schedule[faculty.id][day]:
                        return False
                    if s in classroom_schedule[classroom.id][day]:
                        return False
                    if (course.department, course.semester) in dept_sem_schedule:
                        if s in dept_sem_schedule[(course.department, course.semester)][day]:
                            return False
            else:
                # Single hour course
                if slot in faculty_schedule[faculty.id][day]:
                    return False
                if slot in classroom_schedule[classroom.id][day]:
                    return False
                if (course.department, course.semester) in dept_sem_schedule:
                    if slot in dept_sem_schedule[(course.department, course.semester)][day]:
                        return False
            return True

        # Try to fill each day compactly per dept-semester
        for course in all_courses:
            best_option = None
            best_score = -1e9

            for slot_info in all_slots:
                day = slot_info[0]
                slot = slot_info[1]
                for classroom in all_classrooms.values():
                    faculty = course.faculty
                    if not faculty:
                        continue

                    if not available(course, faculty, classroom, day, slot):
                        continue

                    # --- Scoring: Compact & Balanced ---
                    score = 0

                    # Compact with other same dept-sem classes that day
                    dept_slots = dept_sem_schedule.get((course.department, course.semester), {}).get(day, [])
                    if dept_slots:
                        # Prefer adjacent slots
                        if any(abs(slot.hour - s.hour) == 1 for s in dept_slots):
                            score += 2000
                        else:
                            score -= 100
                    else:
                        score += 300

                    # Faculty idle time minimization
                    fac_slots = faculty_schedule[faculty.id][day]
                    if fac_slots:
                        if any(abs(slot.hour - s.hour) == 1 for s in fac_slots):
                            score += 1200
                        else:
                            score -= 200
                    else:
                        score += 500

                    # Penalize overloading the same day (no penalty for Monday to pack more)
                    num_classes_today = len(dept_slots)
                    if day != 'Monday':
                        score -= num_classes_today * 150

                    # Prefer earlier days in week (increased to heavily favor Monday)
                    score += (5 - self.days.index(day)) * 300

                    # Extra bonus for multi-hour courses on Monday to pack them there
                    if day == 'Monday' and course.duration > 1:
                        score += 1000

                    # Further reduce preference for earlier slots to allow packing later slots on Monday
                    if day == 'Monday' and course.duration > 1:
                        # Prefer later slots for multi-hour courses on Monday
                        score += slot.hour * 50
                    else:
                        score += (17 - slot.hour) * 50

                    if score > best_score:
                        best_score = score
                        best_option = (day, slot, classroom, faculty)

            if best_option:
                day, slot, classroom, faculty = best_option
                timetable.append((course, day, slot, classroom, faculty))

                # Mark all consecutive slots as occupied for multi-hour courses
                duration = course.duration
                for i in range(duration):
                    next_slot = (datetime.combine(datetime.today(), slot) + timedelta(hours=i)).time()
                    faculty_schedule[faculty.id][day].append(next_slot)
                    classroom_schedule[classroom.id][day].append(next_slot)
                    if (course.department, course.semester) not in dept_sem_schedule:
                        dept_sem_schedule[(course.department, course.semester)] = {d: [] for d in self.days}
                    dept_sem_schedule[(course.department, course.semester)][day].append(next_slot)

        # Get next global generation number
        max_gen = db.session.query(db.func.max(Timetable.generation)).scalar() or 0
        new_gen = max_gen + 1

        # Save results
        for course, day, slot, classroom, faculty in timetable:
            new_entry = Timetable(
                department=course.department,
                semester=course.semester,
                course_id=course.id,
                faculty_id=faculty.id,
                classroom_id=classroom.id,
                day=day,
                start_time=slot,
                end_time=(datetime.combine(datetime.today(), slot) + timedelta(hours=course.duration)).time(),
                generation=new_gen
            )
            db.session.add(new_entry)

        db.session.commit()
        self.current_gen = new_gen
        self._optimize_schedule()
        return True, "Timetable generated successfully"

    def _optimize_schedule(self):
        """Shift courses to earlier slots if possible without conflicts"""
        # Get all current timetables for this generation
        current_tts = Timetable.query.filter_by(generation=self.current_gen).all()
        for tt in current_tts:
            self._shift_course_up(tt)

    def _shift_course_up(self, tt):
        """Attempt to shift a timetable entry to an earlier slot"""
        # Find earlier slots on the same day
        current_slot = tt.start_time
        day = tt.day
        duration = tt.course.duration
        earlier_slots = [s for s in self.slots if s < current_slot]

        for new_slot in sorted(earlier_slots, reverse=True):  # Try earliest first
            # Check if available for the entire block
            faculty_schedule = {f.id: {d: [] for d in self.days} for f in self.faculties.values()}
            classroom_schedule = {c.id: {d: [] for d in self.days} for c in self.classrooms.values()}
            dept_sem_schedule = {}

            # Populate schedules from current generation
            all_tts = Timetable.query.filter_by(generation=self.current_gen).all()
            for t in all_tts:
                if t.id != tt.id:  # Exclude current
                    faculty_schedule[t.faculty_id][t.day].append(t.start_time)
                    classroom_schedule[t.classroom_id][t.day].append(t.start_time)
                    key = (t.department, t.semester)
                    if key not in dept_sem_schedule:
                        dept_sem_schedule[key] = {d: [] for d in self.days}
                    dept_sem_schedule[key][t.day].append(t.start_time)

            # Check availability for multi-hour block
            available = True
            slots_needed = []
            for i in range(duration):
                next_slot = (datetime.combine(datetime.today(), new_slot) + timedelta(hours=i)).time()
                if next_slot not in self.slots:
                    available = False
                    break
                slots_needed.append(next_slot)

            if not available:
                continue

            for s in slots_needed:
                if s in faculty_schedule[tt.faculty_id][day]:
                    available = False
                    break
                if s in classroom_schedule[tt.classroom_id][day]:
                    available = False
                    break
                key = (tt.department, tt.semester)
                if key in dept_sem_schedule and s in dept_sem_schedule[key][day]:
                    available = False
                    break

            if available:
                # Shift the entire block
                tt.start_time = new_slot
                tt.end_time = (datetime.combine(datetime.today(), new_slot) + timedelta(hours=duration)).time()
                db.session.commit()
                break  # Shifted to earliest possible

# Routes
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/auth', methods=['GET', 'POST'])
def auth():
    if 'user_id' in session:
        return redirect(url_for('dashboard'))
    login_form = LoginForm(prefix='login')
    register_form = RegistrationForm(prefix='register')

    if login_form.submit.data and login_form.validate_on_submit():
        user = User.query.filter_by(email=login_form.email.data).first()
        if user and user.check_password(login_form.password.data):
            session['user_id'] = user.id
            session['user_name'] = user.full_name
            session['user_role'] = user.role
            flash('Logged in successfully.', 'success')
            return redirect(url_for('dashboard'))
        else:
            flash('Invalid email or password.', 'danger')

    if register_form.submit.data and register_form.validate_on_submit():
        existing_user = User.query.filter_by(email=register_form.email.data).first()
        if existing_user:
            flash('Email already registered. Please login.', 'danger')
            return redirect(url_for('auth'))
        user = User(full_name=register_form.full_name.data, email=register_form.email.data, role=register_form.role.data)
        if register_form.role.data == 'student':
            if not register_form.year.data or not register_form.semester.data:
                flash('Year and Semester are required for students.', 'danger')
                return render_template('auth.html', login_form=login_form, register_form=register_form)
            user.department = register_form.department.data
            user.year = register_form.year.data
            user.semester = register_form.semester.data
        elif register_form.role.data == 'faculty':
            user.department = register_form.department.data
        # For admin, no department, year, semester
        user.set_password(register_form.password.data)
        db.session.add(user)
        if user.role == 'faculty':
            faculty = Faculty(name=user.full_name, availability='Mon,Wed,Fri 10:00-17:00', max_load=5,
                              department=user.department)
            db.session.add(faculty)
        db.session.commit()
        flash('Account created successfully. Please login.', 'success')
        return redirect(url_for('auth'))

    return render_template('auth.html', login_form=login_form, register_form=register_form)

@app.route('/logout')
def logout():
    session.clear()
    flash('You have been logged out.', 'info')
    return redirect(url_for('index'))

@app.route('/profile', methods=['GET', 'POST'])
def profile():
    if 'user_id' not in session:
        flash('Please login to access your profile.', 'warning')
        return redirect(url_for('auth'))
    user = db.session.get(User, session['user_id'])
    if not user:
        session.clear()
        flash('Session expired. Please login again.', 'warning')
        return redirect(url_for('auth'))
    form = ProfileForm(obj=user)
    if form.validate_on_submit():
        user.full_name = form.full_name.data
        if form.password.data:
            user.set_password(form.password.data)
        if user.role == 'student':
            user.department = form.department.data or user.department
            if form.year.data is not None:
                user.year = form.year.data
            if form.semester.data is not None:
                user.semester = form.semester.data
        elif user.role == 'faculty':
            faculty = Faculty.query.filter_by(name=user.full_name).first()
            if faculty:
                faculty.department = form.department.data or faculty.department
        db.session.commit()
        flash('Profile updated successfully.', 'success')
        return redirect(url_for('profile'))
    return render_template('profile.html', form=form, user=user)

@app.route('/delete_account', methods=['POST'])
def delete_account():
    if 'user_id' not in session:
        flash('Please login to delete your account.', 'warning')
        return redirect(url_for('auth'))

    user = db.session.get(User, session['user_id'])
    if not user:
        flash('User not found.', 'danger')
        return redirect(url_for('dashboard'))
    
    # Optional: Check CSRF token if provided
    if request.form.get('csrf_token'):
        # Basic CSRF check (in production, use proper CSRF protection)
        pass  # Assuming token is validated via template
    
    try:
        if user.role == 'faculty':
            # Delete corresponding Faculty entry
            faculty = Faculty.query.filter_by(name=user.full_name).first()
            if faculty:
                # Nullify references
                Course.query.filter_by(faculty_id=faculty.id).update({'faculty_id': None})
                Timetable.query.filter_by(faculty_id=faculty.id).update({'faculty_id': None})
                db.session.delete(faculty)
        
        db.session.delete(user)
        db.session.commit()
        flash('Your account has been deleted successfully.', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error deleting account: {str(e)}', 'danger')
        return redirect(url_for('profile'))
    
    session.clear()
    return redirect(url_for('index'))

@app.route('/add_course', methods=['GET', 'POST'])
def add_course():
    if 'user_id' not in session:
        flash('Please login to access this page.', 'warning')
        return redirect(url_for('auth'))
    
    user_role = session.get('user_role')
    if user_role == 'student':
        flash('Access denied. Students cannot add courses. Use the Enroll button to join available courses.', 'danger')
        return redirect(url_for('dashboard'))
    
    form = CourseForm()
    form.classroom_id.choices = [(c.id, c.name) for c in Classroom.query.all()]
    
    if user_role == 'faculty':
        # For faculty, only allow their own faculty
        faculty = Faculty.query.filter_by(name=session['user_name']).first()
        if not faculty:
            flash('Faculty profile not found. Contact admin.', 'warning')
            return redirect(url_for('dashboard'))
        form.faculty_id.choices = [(faculty.id, faculty.name)]
        form.faculty_id.data = faculty.id
    else:
        # For admin or others, all faculties
        form.faculty_id.choices = [(f.id, f.name) for f in Faculty.query.all()]
    
    if form.validate_on_submit():
        course = Course(name=form.name.data, faculty_id=form.faculty_id.data,
                        classroom_id=form.classroom_id.data, duration=form.duration.data,
                        department=form.department.data, year=form.year.data, semester=form.semester.data)
        db.session.add(course)
        db.session.commit()
        flash('Course added successfully!')
        return redirect(url_for('dashboard'))
    return render_template('add_course.html', form=form, user_role=user_role)

@app.route('/edit_course/<int:course_id>', methods=['GET', 'POST'])
def edit_course(course_id):
    if 'user_id' not in session:
        flash('Please login to access this page.', 'warning')
        return redirect(url_for('auth'))
    
    user_role = session.get('user_role')
    if user_role == 'student':
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))
    
    course = Course.query.get_or_404(course_id)
    if user_role == 'faculty':
        faculty = Faculty.query.filter_by(name=session['user_name']).first()
        if not faculty or course.faculty_id != faculty.id:
            flash('Access denied.', 'danger')
            return redirect(url_for('dashboard'))
    
    form = CourseForm(obj=course)
    form.classroom_id.choices = [(c.id, c.name) for c in Classroom.query.all()]
    
    if user_role == 'faculty':
        faculty = Faculty.query.filter_by(name=session['user_name']).first()
        if faculty:
            form.faculty_id.choices = [(faculty.id, faculty.name)]
            form.faculty_id.data = faculty.id
    else:
        form.faculty_id.choices = [(f.id, f.name) for f in Faculty.query.all()]
    
    if form.validate_on_submit():
        course.name = form.name.data
        course.faculty_id = form.faculty_id.data
        course.classroom_id = form.classroom_id.data
        course.duration = form.duration.data
        course.department = form.department.data
        course.year = form.year.data
        course.semester = form.semester.data
        db.session.commit()
        # Update timetable semester for this course
        Timetable.query.filter_by(course_id=course.id).update({'semester': course.semester})
        db.session.commit()
        flash('Course updated successfully!', 'success')
        return redirect(url_for('dashboard'))
    return render_template('edit_course.html', form=form, user_role=user_role)

@app.route('/edit_classroom/<int:classroom_id>', methods=['GET', 'POST'])
def edit_classroom(classroom_id):
    if 'user_id' not in session:
        flash('Please login to access this page.', 'warning')
        return redirect(url_for('auth'))
    
    user_role = session.get('user_role')
    if user_role == 'student':
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))
    
    classroom = Classroom.query.get_or_404(classroom_id)
    
    form = ClassroomForm(obj=classroom)
    
    if form.validate_on_submit():
        classroom.name = form.name.data
        classroom.capacity = form.capacity.data
        classroom.type = form.type.data
        db.session.commit()
        flash('Classroom updated successfully!', 'success')
        return redirect(url_for('dashboard'))
    return render_template('edit_classroom.html', form=form, user_role=user_role)

@app.route('/enroll_course', methods=['GET', 'POST'])
def enroll_course():
    if 'user_id' not in session:
        flash('Please login to access this page.', 'warning')
        return redirect(url_for('auth'))
    
    user_role = session.get('user_role')
    if user_role != 'student':
        flash('Access denied. Only students can enroll in courses.', 'danger')
        return redirect(url_for('dashboard'))
    
    user = User.query.get(session['user_id'])
    available_courses = [c for c in Course.query.all() if c not in user.enrolled_courses and c.department == user.department and c.year == user.year and c.semester == user.semester]

    if request.method == 'POST' and 'course_id' in request.form:
        course_id = int(request.form['course_id'])
        course = Course.query.get(course_id)
        if course and course in available_courses:
            user.enrolled_courses.append(course)
            db.session.commit()
            flash(f'Successfully enrolled in {course.name}!', 'success')
            return redirect(url_for('dashboard'))
        else:
            flash('Invalid course selection.', 'danger')
            return redirect(url_for('enroll_course'))
    
    form = EnrollmentForm()
    form.course_id.choices = [(c.id, f"{c.name} (Faculty: {c.faculty.name if c.faculty else 'Unassigned'}, Classroom: {c.classroom.name if c.classroom else 'Unassigned'})") for c in available_courses]
    
    if form.validate_on_submit():
        course = Course.query.get(form.course_id.data)
        if course:
            user.enrolled_courses.append(course)
            db.session.commit()
            flash(f'Successfully enrolled in {course.name}!', 'success')
        else:
            flash('Invalid course selection.', 'danger')
        return redirect(url_for('dashboard'))
    
    return render_template('enroll_course.html', form=form)

@app.route('/unenroll_course/<int:course_id>', methods=['POST'])
def unenroll_course(course_id):
    if 'user_id' not in session:
        flash('Please login to access this page.', 'warning')
        return redirect(url_for('auth'))
    
    user_role = session.get('user_role')
    if user_role != 'student':
        flash('Access denied. Only students can unenroll from courses.', 'danger')
        return redirect(url_for('dashboard'))
    
    user = User.query.get(session['user_id'])
    course = Course.query.get_or_404(course_id)
    
    if course in user.enrolled_courses:
        user.enrolled_courses.remove(course)
        db.session.commit()
        flash(f'Unenrolled from {course.name}!', 'success')
    else:
        flash('Not enrolled in this course.', 'warning')
    
    return redirect(url_for('dashboard'))

@app.route('/delete_course/<int:course_id>', methods=['POST'])
def delete_course(course_id):
    if 'user_id' not in session:
        flash('Please login to access this page.', 'warning')
        return redirect(url_for('auth'))
    user_role = session.get('user_role')
    user_name = session.get('user_name')
    course = Course.query.get_or_404(course_id)
    if user_role == 'faculty':
        faculty = Faculty.query.filter_by(name=user_name).first()
        if not faculty or course.faculty_id != faculty.id:
            flash('Access denied.', 'danger')
            return redirect(url_for('dashboard'))
    elif user_role not in ['admin', 'student']:
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))
    try:
        db.session.delete(course)
        db.session.commit()
        flash('Course deleted successfully!', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error deleting course: {str(e)}', 'danger')
    return redirect(url_for('dashboard'))

@app.route('/delete_faculty/<int:faculty_id>', methods=['POST'])
def delete_faculty(faculty_id):
    if 'user_id' not in session:
        flash('Please login to access this page.', 'warning')
        return redirect(url_for('auth'))
    user_role = session.get('user_role')
    if user_role != 'admin':
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))
    faculty = Faculty.query.get_or_404(faculty_id)
    try:
        # Nullify references in related models
        Course.query.filter_by(faculty_id=faculty.id).update({'faculty_id': None})
        Timetable.query.filter_by(faculty_id=faculty.id).update({'faculty_id': None})
        db.session.delete(faculty)
        db.session.commit()
        flash('Faculty deleted successfully!', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error deleting faculty: {str(e)}', 'danger')
    return redirect(url_for('dashboard'))

@app.route('/delete_classroom/<int:classroom_id>', methods=['POST'])
def delete_classroom(classroom_id):
    if 'user_id' not in session:
        flash('Please login to access this page.', 'warning')
        return redirect(url_for('auth'))
    user_role = session.get('user_role')
    if user_role not in ['faculty', 'admin']:
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))
    classroom = Classroom.query.get_or_404(classroom_id)
    try:
        db.session.delete(classroom)
        db.session.commit()
        flash('Classroom deleted successfully!', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error deleting classroom: {str(e)}', 'danger')
    return redirect(url_for('dashboard'))

@app.route('/delete_timetable/<int:tt_id>', methods=['POST'])
def delete_timetable(tt_id):
    if 'user_id' not in session:
        flash('Please login to access this page.', 'warning')
        return redirect(url_for('auth'))
    user_role = session.get('user_role')
    if user_role != 'admin':
        flash('Access denied.', 'danger')
        return redirect(url_for('generate'))
    tt = Timetable.query.get_or_404(tt_id)
    try:
        db.session.delete(tt)
        db.session.commit()
        flash('Timetable entry deleted successfully!', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error deleting timetable entry: {str(e)}', 'danger')
    return redirect(url_for('generate'))

@app.route('/add_faculty', methods=['GET', 'POST'])
def add_faculty():
    if 'user_id' not in session:
        flash('Please login to access this page.', 'warning')
        return redirect(url_for('auth'))
    user_role = session.get('user_role')
    if user_role not in ['admin']:
        flash('Access denied. Only admins can add faculty.', 'danger')
        return redirect(url_for('dashboard'))
    form = FacultyForm()
    if form.validate_on_submit():
        faculty = Faculty(name=form.name.data, availability=form.availability.data, max_load=form.max_load.data,
                          department=form.department.data, year=form.year.data, semester=form.semester.data)
        db.session.add(faculty)
        db.session.commit()
        flash('Faculty added successfully!')
        return redirect(url_for('dashboard'))
    return render_template('add_faculty.html', form=form)

@app.route('/add_classroom', methods=['GET', 'POST'])
def add_classroom():
    if 'user_id' not in session:
        flash('Please login to access this page.', 'warning')
        return redirect(url_for('auth'))
    user_role = session.get('user_role')
    if user_role == 'student':
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))
    form = ClassroomForm()
    if form.validate_on_submit():
        classroom = Classroom(name=form.name.data, capacity=form.capacity.data, type=form.type.data)
        db.session.add(classroom)
        db.session.commit()
        flash('Classroom added successfully!')
        return redirect(url_for('dashboard'))
    return render_template('add_classroom.html', form=form)

@app.route('/generate_timetable', methods=['GET', 'POST'])
def generate():
    if 'user_id' not in session:
        flash('Please login to access this page.', 'warning')
        return redirect(url_for('auth'))

    user_role = session.get('user_role')

    if request.method == 'POST':
        if user_role != 'admin':
            flash('Access denied. Only admins can generate timetables.', 'danger')
            return redirect(url_for('dashboard'))

        # Get all schedulable courses across all departments and semesters
        courses_to_schedule = Course.query.all()
        schedulable_courses = [c for c in courses_to_schedule if c.faculty_id and c.classroom_id]

        if not schedulable_courses:
            flash('No courses with assigned faculty and classroom found.', 'warning')
        else:
            scheduler = ConflictFreeScheduler(schedulable_courses)
            success, message = scheduler.generate()
            if success:
                flash(message)
            else:
                flash(message, 'warning')

    # Get all timetables
    timetables = Timetable.query.options(
        db.joinedload(Timetable.course),
        db.joinedload(Timetable.faculty_obj),
        db.joinedload(Timetable.classroom_obj)
    ).all()

    # Group timetables by department and semester
    timetable_groups = {}
    for tt in timetables:
        key = (tt.department, tt.semester)
        if key not in timetable_groups:
            timetable_groups[key] = []
        timetable_groups[key].append(tt)

    if user_role == 'student':
        user = User.query.get(session['user_id'])
        courses = user.enrolled_courses
        faculties = Faculty.query.filter_by(department=user.department).all()
        # Filter timetable_groups to only include groups where courses are enrolled
        filtered_groups = {}
        for key, tts in timetable_groups.items():
            filtered_tts = [tt for tt in tts if tt.course in courses]
            if filtered_tts:
                filtered_groups[key] = filtered_tts
        timetable_groups = filtered_groups
    elif user_role == 'faculty':
        # Filter timetable_groups to only the faculty's department
        faculty = Faculty.query.filter_by(name=session['user_name']).first()
        if faculty:
            filtered_groups = {}
            for key, tts in timetable_groups.items():
                if key[0] == faculty.department:
                    filtered_groups[key] = tts
            timetable_groups = filtered_groups
        else:
            timetable_groups = {}
    # For admin, show all groups

    courses = Course.query.all()
    faculties = Faculty.query.all()
    classrooms = Classroom.query.all()
    return render_template('generate_timetable.html', timetables=timetables, courses=courses, faculties=faculties, classrooms=classrooms, user_role=user_role, timetable_groups=timetable_groups)

@app.route('/dashboard')
def dashboard():
    if 'user_id' not in session:
        flash('Please login to access the dashboard.', 'warning')
        return redirect(url_for('auth'))

    user = User.query.get(session['user_id'])
    if not user:
        session.clear()
        flash('Session expired. Please login again.', 'warning')
        return redirect(url_for('auth'))

    user_role = session.get('user_role')
    user_name = session.get('user_name')

    all_timetables = Timetable.query.options(
        db.joinedload(Timetable.course),
        db.joinedload(Timetable.faculty_obj),
        db.joinedload(Timetable.classroom_obj)
    ).all()
    all_courses = Course.query.all()
    all_faculties = Faculty.query.all()
    all_classrooms = Classroom.query.all()

    timetables = all_timetables
    courses = all_courses
    faculties = all_faculties
    classrooms = all_classrooms
    current_faculty_id = None
    enrolled_courses = []
    available_courses = []
    form = None
    timetable_groups = {}  # For admin: group by (department, semester)

    if user_role == 'faculty':
        # Find the faculty matching the user's name
        faculty = Faculty.query.filter_by(name=user_name).first()
        if faculty:
            current_faculty_id = faculty.id
            timetables = [tt for tt in all_timetables if tt.faculty_id == faculty.id]
            # For faculty, show only their courses
            courses = [c for c in all_courses if c.faculty_id == faculty.id]
            # Limit faculties to their department for modal
            faculties = Faculty.query.filter_by(department=faculty.department).all()
            # Show all classrooms
            classrooms = all_classrooms
        else:
            flash('Faculty profile not found. Contact admin.', 'warning')
            faculties = []
            classrooms = []

    elif user_role == 'student':
        user = User.query.get(session['user_id'])
        enrolled_courses = user.enrolled_courses
        # Filter courses by student's department, year, semester
        matching_courses = [c for c in all_courses if c.department == user.department and c.year == user.year and c.semester == user.semester]
        available_courses = [c for c in matching_courses if c not in enrolled_courses]
        form = EnrollmentForm()
        form.course_id.choices = [(c.id, f"{c.name} (Faculty: {c.faculty.name if c.faculty else 'Unassigned'}, Classroom: {c.classroom.name if c.classroom else 'Unassigned'})") for c in available_courses]
        # For students, show matching courses and timetables for enrolled courses
        courses = matching_courses
        timetables = [tt for tt in all_timetables if tt.course in enrolled_courses] if enrolled_courses else []
        faculties = Faculty.query.filter_by(department=user.department).all()
        classrooms = all_classrooms
        current_faculty_id = None

    elif user_role == 'admin':
        # Group timetables by department and semester
        for tt in all_timetables:
            key = (tt.department, tt.semester)
            if key not in timetable_groups:
                timetable_groups[key] = []
            timetable_groups[key].append(tt)

    enrolled_courses = enrolled_courses if user_role == 'student' else []

    return render_template('dashboard.html', timetables=timetables, courses=courses, faculties=faculties, classrooms=classrooms, user_role=user_role, current_faculty_id=current_faculty_id, enrolled_courses=enrolled_courses, available_courses=available_courses if user_role == 'student' else None, form=form if user_role == 'student' else None, timetable_groups=timetable_groups)

@app.route('/export_pdf')
def export_pdf():
    department = request.args.get('department')
    semester = request.args.get('semester')
    if department and semester:
        try:
            semester = int(semester)
            timetables = Timetable.query.filter_by(department=department, semester=semester).all()
        except ValueError:
            timetables = Timetable.query.all()
    else:
        timetables = Timetable.query.all()
    buffer = io.BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    p.drawString(100, 750, f"Timetable - {department.upper() if department else 'All'} Semester {semester if semester else 'All'}")
    y = 700
    line_height = 15
    for tt in timetables:
        p.drawString(100, y, f"Day: {tt.day}")
        y -= line_height
        p.drawString(100, y, f"Time: {tt.start_time.strftime('%I:%M %p')} - {tt.end_time.strftime('%I:%M %p')}")
        y -= line_height
        p.drawString(100, y, f"Course: {tt.course.name}")
        y -= line_height
        p.drawString(100, y, f"Faculty: {tt.faculty_obj.name}")
        y -= line_height
        p.drawString(100, y, f"Classroom: {tt.classroom_obj.name}")
        y -= line_height * 2  # extra space between entries
        if y < 100:
            p.showPage()
            y = 700
    p.showPage()
    p.save()
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name=f'timetable_{department}_{semester}.pdf' if department and semester else 'timetable.pdf', mimetype='application/pdf')

@app.route('/export_doc')
def export_doc():
    department = request.args.get('department')
    semester = request.args.get('semester')
    if department and semester:
        try:
            semester = int(semester)
            timetables = Timetable.query.filter_by(department=department, semester=semester).all()
        except ValueError:
            timetables = Timetable.query.all()
    else:
        timetables = Timetable.query.all()
    doc = Document()
    doc.add_heading(f'Timetable - {department.upper() if department else "All"} Semester {semester if semester else "All"}', 0)
    for tt in timetables:
        doc.add_paragraph(f"Day: {tt.day}\nTime: {tt.start_time.strftime('%I:%M %p')} - {tt.end_time.strftime('%I:%M %p')}\nCourse: {tt.course.name}\nFaculty: {tt.faculty_obj.name}\nClassroom: {tt.classroom_obj.name}")
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name=f'timetable_{department}_{semester}.docx' if department and semester else 'timetable.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.route('/validate_timetables')
def validate_timetables():
    if 'user_id' not in session:
        flash('Please login to access this page.', 'warning')
        return redirect(url_for('auth'))

    user_role = session.get('user_role')
    if user_role not in ['admin', 'faculty']:
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))

    # Detect conflicts across all timetables
    conflicts = []
    all_timetables = Timetable.query.options(
        db.joinedload(Timetable.course),
        db.joinedload(Timetable.faculty_obj),
        db.joinedload(Timetable.classroom_obj)
    ).all()

    # Group by day and time slot
    schedule = {}
    for tt in all_timetables:
        if not tt.day or not tt.start_time or not tt.end_time:
            continue
        key = (tt.day, tt.start_time, tt.end_time)
        if key not in schedule:
            schedule[key] = []
        schedule[key].append(tt)

    for (day, start, end), tts in schedule.items():
        if len(tts) > 1:
            # Check for conflicts
            faculties = set()
            classrooms = set()
            dept_sems = set()
            conflicting_tts = []
            for tt in tts:
                if tt.faculty_id in faculties:
                    conflicts.append({
                        'type': 'Faculty Conflict',
                        'severity': 'critical',
                        'description': f"Faculty {tt.faculty_obj.name} is double-booked on {day} at {start.strftime('%I:%M %p')}",
                        'details': [f"{t.course.name} ({t.faculty_obj.name})" for t in tts],
                        'timetable_ids': [t.id for t in tts]
                    })
                if tt.classroom_id in classrooms:
                    conflicts.append({
                        'type': 'Classroom Conflict',
                        'severity': 'critical',
                        'description': f"Classroom {tt.classroom_obj.name} is double-booked on {day} at {start.strftime('%I:%M %p')}",
                        'details': [f"{t.course.name} ({t.classroom_obj.name})" for t in tts],
                        'timetable_ids': [t.id for t in tts]
                    })
                if (tt.department, tt.semester) in dept_sems:
                    conflicts.append({
                        'type': 'Student Group Conflict',
                        'severity': 'high',
                        'description': f"Students in {tt.department.upper()} Semester {tt.semester} have overlapping classes on {day} at {start.strftime('%I:%M %p')}",
                        'details': [f"{t.course.name}" for t in tts],
                        'timetable_ids': [t.id for t in tts]
                    })
                faculties.add(tt.faculty_id)
                classrooms.add(tt.classroom_id)
                dept_sems.add((tt.department, tt.semester))

    return render_template('validate_timetables.html', conflicts=conflicts)

@app.route('/fix_conflicts', methods=['POST'])
def fix_conflicts():
    if 'user_id' not in session:
        flash('Please login to access this page.', 'warning')
        return redirect(url_for('auth'))

    user_role = session.get('user_role')
    if user_role != 'admin':
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))

    # Find departments and semesters with conflicts and regenerate
    # For simplicity, regenerate all existing timetables
    existing_dept_sem = db.session.query(Timetable.department, Timetable.semester).distinct().all()
    for dept, sem in existing_dept_sem:
        courses = Course.query.filter_by(department=dept, semester=sem).all()
        schedulable = [c for c in courses if c.faculty_id and c.classroom_id]
        if schedulable:
            scheduler = ConflictFreeScheduler(schedulable)
            success, message = scheduler.generate()
            if not success:
                flash(f"Failed to regenerate for {dept.upper()} Sem {sem}: {message}", 'warning')

    flash('Attempted to fix conflicts by regenerating affected timetables.', 'info')
    return redirect(url_for('validate_timetables'))

@app.route('/clear_all_timetables', methods=['POST'])
def clear_all_timetables():
    if 'user_id' not in session:
        flash('Please login to access this page.', 'warning')
        return redirect(url_for('auth'))

    user_role = session.get('user_role')
    if user_role != 'admin':
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))

    Timetable.query.delete()
    db.session.commit()
    flash('All timetables have been cleared.', 'success')
    return redirect(url_for('validate_timetables'))



@app.route('/clock.jpg')
def clock_image():
    return send_file('clock.jpg', mimetype='image/jpeg')

# Vercel deployment - create tables on startup
with app.app_context():
    db.create_all()

from vercel_wsgi import handle

def handler(event, context):
    return handle(app, event, context)




